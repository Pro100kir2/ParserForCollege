from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from typing import List
from src.domain.interfaces import Exporter
from src.domain.models import ParseResult, SectionType
from src.utils.logger import get_logger
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.oxml import OxmlElement

logger = get_logger(__name__)

DARK_BLUE = RGBColor(31, 78, 121)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(120, 120, 120)

TABLE_HEADER = "D9EAF7"

class WordExporter(Exporter):
    def __init__(self):
        pass

    def _set_cell_background(self, cell, color):
        tc_pr = cell._tc.get_or_add_tcPr()

        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color)

        tc_pr.append(shd)

    def _apply_professional_styling(self, doc: Document):
        """
        Университетский стиль документа
        """

        section = doc.sections[0]

        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

        normal = doc.styles["Normal"]

        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)
        normal.font.color.rgb = BLACK

        normal.paragraph_format.first_line_indent = Cm(1.25)
        normal.paragraph_format.line_spacing = 1.15
        normal.paragraph_format.space_after = Pt(12)

        h1 = doc.styles["Heading 1"]

        h1.font.name = "Times New Roman"
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = DARK_BLUE

        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(12)

        h2 = doc.styles["Heading 2"]

        h2.font.name = "Times New Roman"
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = DARK_BLUE

        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(8)

        title = doc.styles["Title"]

        title.font.name = "Times New Roman"
        title.font.size = Pt(24)
        title.font.bold = True
        title.font.color.rgb = DARK_BLUE

    def _add_cover_page(self, doc: Document, result: ParseResult):

        title = doc.add_heading("ОТЧЕТ ПО ОБРАБОТКЕ РПД", 0)

        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()

        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = subtitle.add_run(
            "Рабочие программы дисциплин"
        )

        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.color.rgb = BLACK

        subtitle.paragraph_format.space_after = Pt(20)

        line = doc.add_paragraph()

        line.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = line.add_run(
            "────────────────────────────────────────────"
        )

        run.font.color.rgb = DARK_BLUE

        doc.add_paragraph()

        stats_table = doc.add_table(
            rows=4,
            cols=2
        )

        stats_table.style = "Table Grid"

        stats_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        header = stats_table.rows[0].cells

        header[0].text = "Метрика"
        header[1].text = "Значение"

        for cell in header:
            self._set_cell_background(
                cell,
                TABLE_HEADER
            )

            paragraph = cell.paragraphs[0]

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            run = paragraph.runs[0]

            run.bold = True
            run.font.color.rgb = BLACK

        stats_table.rows[1].cells[0].text = (
            "Обработано файлов"
        )

        stats_table.rows[1].cells[1].text = (
            str(result.processed_files)
        )

        stats_table.rows[2].cells[0].text = (
            "Ошибок"
        )

        stats_table.rows[2].cells[1].text = (
            str(result.failed_files)
        )

        stats_table.rows[3].cells[0].text = (
            "Версия экстрактора"
        )

        stats_table.rows[3].cells[1].text = (
            result.extractor_version
        )

        for row in stats_table.rows[1:]:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER
                    )

                    for run in paragraph.runs:
                        run.font.color.rgb = BLACK

        doc.add_page_break()

    def export(self, result: ParseResult, config):
        """Export results to Word document with professional styling"""
        doc = Document()

        # Apply professional styling
        self._apply_professional_styling(doc)

        # Add cover page
        self._add_cover_page(doc, result)

        # Add executive summary
        doc.add_heading('Исполнительное резюме', level=1)

        summary = doc.add_paragraph()
        summary.add_run('В ходе обработки было проанализировано ').bold = True
        summary.add_run(f'{result.processed_files} ').bold = True
        summary.add_run('документов. ')

        if result.failed_files > 0:
            summary.add_run(f'Обнаружено ').bold = True
            summary.add_run(f'{result.failed_files} ').bold = True
            summary.add_run('ошибок. ')
        else:
            summary.add_run('Ошибок не обнаружено. ')

        summary.add_run(f'Извлечено ').bold = True
        summary.add_run(f'{len(result.main_literature)} ').bold = True
        summary.add_run('записей основной литературы, ')
        summary.add_run(f'{len(result.additional_literature)} ').bold = True
        summary.add_run('записей дополнительной литературы и ')
        summary.add_run(f'{len(result.material_resources)} ').bold = True
        summary.add_run('записей материально-технического обеспечения.')

        # Add issues section if any
        if result.failed_file_list or result.low_confidence_files:
            doc.add_heading('Выявленные проблемы', level=1)

            if result.failed_file_list:
                doc.add_heading('Файлы с ошибками обработки', level=2)
                for filename in result.failed_file_list:
                    p = doc.add_paragraph(filename, style='List Bullet')
                    p.runs[0].font.color.rgb = RGBColor(192, 0, 0)  # Red

            if result.low_confidence_files:
                doc.add_heading('Файлы с низкой уверенностью извлечения', level=2)
                p = doc.add_paragraph('Следующие файлы требуют ручной проверки:')
                p.runs[0].font.italic = True
                for filename in result.low_confidence_files:
                    p = doc.add_paragraph(filename, style='List Bullet')
                    p.runs[0].font.color.rgb = RGBColor(255, 140, 0)  # Orange

        # Add main literature section
        if result.main_literature:
            doc.add_page_break()
            doc.add_heading('Основная литература', level=1)
            p = doc.add_paragraph(f'Всего извлечено записей: {len(result.main_literature)}')
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(89, 89, 89)

            for i, item in enumerate(result.main_literature, 1):
                p = doc.add_paragraph(f'{i}. {item}')
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(12)

        # Add additional literature section
        if result.additional_literature:
            doc.add_page_break()
            doc.add_heading('Дополнительная литература', level=1)
            p = doc.add_paragraph(f'Всего извлечено записей: {len(result.additional_literature)}')
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(89, 89, 89)

            for i, item in enumerate(result.additional_literature, 1):
                p = doc.add_paragraph(f'{i}. {item}')
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(12)

        # Add material resources section
        if result.material_resources:
            doc.add_page_break()
            doc.add_heading('Материально-техническое обеспечение', level=1)
            p = doc.add_paragraph(f'Всего извлечено записей: {len(result.material_resources)}')
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(89, 89, 89)

            for item in result.material_resources:
                p = doc.add_paragraph(f'• {item}')
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(12)

        # Add footer with page numbers
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add page number field
        from docx.oxml.shared import OxmlElement, qn
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        footer_para._p.append(fldChar1)
        footer_para._p.append(instrText)
        footer_para._p.append(fldChar2)

        # Save document
        output_path = config.output_dir
        filename = config.combined_filename.replace('.xlsx', '.docx')
        full_path = f"{output_path}/{filename}"

        doc.save(full_path)
        logger.info(f"Created professional Word document: {full_path}")

        return [full_path]

    def export_multiple(self, result: ParseResult, config):
        """Export results to multiple Word documents with professional styling"""
        output_files = []
        output_path = config.output_dir

        # Main literature
        if result.main_literature:
            doc = Document()
            self._apply_professional_styling(doc)

            # Title
            title = doc.add_heading('Основная литература', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Summary
            p = doc.add_paragraph(f'Всего извлечено записей: {len(result.main_literature)}')
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(89, 89, 89)
            p.paragraph_format.space_after = Pt(12)

            # Entries
            for i, item in enumerate(result.main_literature, 1):
                p = doc.add_paragraph(f'{i}. {item}')
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(12)

            # Footer with page numbers
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.runs[0].font.size = Pt(9)
            footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # Add page number field
            from docx.oxml.shared import OxmlElement, qn
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            footer_para._p.append(fldChar1)
            footer_para._p.append(instrText)
            footer_para._p.append(fldChar2)

            filename = config.filename_main.replace('.xlsx', '.docx')
            full_path = f"{output_path}/{filename}"
            doc.save(full_path)
            output_files.append(full_path)
            logger.info(f"Created: {full_path}")

        # Additional literature
        if result.additional_literature:
            doc = Document()
            self._apply_professional_styling(doc)

            title = doc.add_heading('Дополнительная литература', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph(f'Всего извлечено записей: {len(result.additional_literature)}')
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(89, 89, 89)
            p.paragraph_format.space_after = Pt(12)

            for i, item in enumerate(result.additional_literature, 1):
                p = doc.add_paragraph(f'{i}. {item}')
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(12)

            # Footer with page numbers
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.runs[0].font.size = Pt(9)
            footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # Add page number field
            from docx.oxml.shared import OxmlElement, qn
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            footer_para._p.append(fldChar1)
            footer_para._p.append(instrText)
            footer_para._p.append(fldChar2)

            filename = config.filename_additional.replace('.xlsx', '.docx')
            full_path = f"{output_path}/{filename}"
            doc.save(full_path)
            output_files.append(full_path)
            logger.info(f"Created: {full_path}")

        # Material resources
        if result.material_resources:
            doc = Document()
            self._apply_professional_styling(doc)

            title = doc.add_heading('Материально-техническое обеспечение', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph(f'Всего извлечено записей: {len(result.material_resources)}')
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(89, 89, 89)
            p.paragraph_format.space_after = Pt(12)

            for item in result.material_resources:
                p = doc.add_paragraph(f'• {item}')
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(12)

            # Footer with page numbers
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.runs[0].font.size = Pt(9)
            footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # Add page number field
            from docx.oxml.shared import OxmlElement, qn
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            footer_para._p.append(fldChar1)
            footer_para._p.append(instrText)
            footer_para._p.append(fldChar2)

            filename = config.filename_material.replace('.xlsx', '.docx')
            full_path = f"{output_path}/{filename}"
            doc.save(full_path)
            output_files.append(full_path)
            logger.info(f"Created: {full_path}")

        logger.info(f"Exported {len(output_files)} file(s) to {output_path}")
        return output_files